"""创建带批注的 Word 文档（python-docx + zipfile 注入 XML）"""
import zipfile, shutil, os
from lxml import etree

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
W14 = 'http://schemas.microsoft.com/office/word/2010/wordml'

def build_comments_xml(entries):
    """entries: [(cid, author, date, text, para_id), ...]"""
    parts = []
    for cid, author, date, text, pid in entries:
        parts.append(f'''  <w:comment w:id="{cid}" w:author="{author}" w:date="{date}" w:initials="{author[0]}">
    <w:p w14:paraId="{pid}" w14:textId="77777777">
      <w:r><w:annotationRef/></w:r>
      <w:r><w:t xml:space="preserve">{text}</w:t></w:r>
    </w:p>
  </w:comment>''')
    return f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:comments xmlns:w="{W}" xmlns:w14="{W14}"
            xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml"
            xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
            mc:Ignorable="w14 w15">
{chr(10).join(parts)}
</w:comments>'''

def inject_markers(xml_bytes, cid, target_text):
    tree = etree.fromstring(xml_bytes)
    for p in tree.iter(f'{{{W}}}p'):
        for r in p.iter(f'{{{W}}}r'):
            t = r.find(f'{{{W}}}t')
            if t is not None and target_text in (t.text or ''):
                idx = list(p).index(r)
                start = etree.Element(f'{{{W}}}commentRangeStart')
                start.set(f'{{{W}}}id', str(cid))
                p.insert(idx, start)
                end = etree.Element(f'{{{W}}}commentRangeEnd')
                end.set(f'{{{W}}}id', str(cid))
                p.insert(idx + 2, end)
                ref_r = etree.Element(f'{{{W}}}r')
                rpr = etree.SubElement(ref_r, f'{{{W}}}rPr')
                etree.SubElement(rpr, f'{{{W}}}rStyle').set(f'{{{W}}}val', 'CommentReference')
                etree.SubElement(ref_r, f'{{{W}}}commentReference').set(f'{{{W}}}id', str(cid))
                p.insert(idx + 3, ref_r)
                break
    return etree.tostring(tree, xml_declaration=True, encoding='UTF-8', standalone=True)

def inject_content_type(xml_bytes):
    tree = etree.fromstring(xml_bytes)
    ns = 'http://schemas.openxmlformats.org/package/2006/content-types'
    for o in tree.findall(f'{{{ns}}}Override'):
        if o.get('PartName') == '/word/comments.xml':
            return etree.tostring(tree, xml_declaration=True, encoding='UTF-8', standalone=True)
    o = etree.SubElement(tree, f'{{{ns}}}Override')
    o.set('PartName', '/word/comments.xml')
    o.set('ContentType', 'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml')
    return etree.tostring(tree, xml_declaration=True, encoding='UTF-8', standalone=True)

def inject_rel(xml_bytes):
    tree = etree.fromstring(xml_bytes)
    ns = 'http://schemas.openxmlformats.org/package/2006/relationships'
    for r in tree.findall(f'{{{ns}}}Relationship'):
        if 'comments' in r.get('Type', '').lower():
            return etree.tostring(tree, xml_declaration=True, encoding='UTF-8', standalone=True)
    r = etree.SubElement(tree, f'{{{ns}}}Relationship')
    r.set('Id', 'rIdComments')
    r.set('Type', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments')
    r.set('Target', 'comments.xml')
    return etree.tostring(tree, xml_declaration=True, encoding='UTF-8', standalone=True)

def create(base_docx, output_docx, comments):
    """
    comments: [(cid, author, date, text, para_id, target_text), ...]
    """
    shutil.copy2(base_docx, output_docx)
    tmp = output_docx + '.tmp'
    entries = [(c[0], c[1], c[2], c[3], c[4]) for c in comments]
    comments_xml = build_comments_xml(entries)

    with zipfile.ZipFile(output_docx, 'r') as zin, \
         zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == 'word/document.xml':
                for cid, _, _, _, _, target in comments:
                    data = inject_markers(data, cid, target)
            elif item.filename == '[Content_Types].xml':
                data = inject_content_type(data)
            elif item.filename == 'word/_rels/document.xml.rels':
                data = inject_rel(data)
            zout.writestr(item, data)
        zout.writestr('word/comments.xml', comments_xml)
    os.replace(tmp, output_docx)
    print(f'✅ Created: {output_docx}')

if __name__ == '__main__':
    from docx import Document
    doc = Document()
    doc.add_paragraph('这是第一段，将被添加批注。')
    doc.add_paragraph('这是第二段，也会被批注。')
    doc.save('base.docx')
    create('base.docx', 'output_with_comments.docx', [
        ('0', 'Claude', '2026-05-13T10:00:00Z', '这是第一条批注', '1D780707', '第一段'),
        ('1', 'Claude', '2026-05-13T10:01:00Z', '这是第二条批注', '1D780708', '第二段'),
    ])
